import io
import os
import json
import uuid
import time
import pathlib
import typing as t
import re

import numpy as np
from fastapi import FastAPI, HTTPException, Body, UploadFile, File, Form
from pydantic import BaseModel, Field
from dotenv import load_dotenv
from docx import Document
from .storage.vector_store import VectorStore, simple_chunk

try:
    from sentence_transformers import SentenceTransformer
    _HAS_ST = True
except Exception:
    _HAS_ST = False

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    _HAS_SK = True
except Exception:
    _HAS_SK = False

try:
    import openai  # type: ignore
    _HAS_OPENAI = True
except Exception:
    _HAS_OPENAI = False
import requests


BASE_DIR = pathlib.Path(__file__).resolve().parent.parent
load_dotenv(dotenv_path=BASE_DIR / ".env")
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(exist_ok=True, parents=True)
STORE_PATH = DATA_DIR / "store.json"
EMB_PATH = DATA_DIR / "embeddings.npy"


class IngestDocument(BaseModel):
    id: t.Optional[str] = None
    text: str
    meta: t.Optional[dict] = Field(default_factory=dict)


class IngestRequest(BaseModel):
    documents: t.List[IngestDocument]


class SearchRequest(BaseModel):
    query: str
    top_k: int = 5


class SearchResult(BaseModel):
    doc_id: str
    chunk_id: str
    score: float
    text: str
    meta: dict


class SearchResponse(BaseModel):
    query: str
    results: t.List[SearchResult]


class AnswerRequest(BaseModel):
    query: str
    top_k: int = 8
    max_tokens: int = 512


class AnswerResponse(BaseModel):
    answer: str
    citations: t.List[dict]
    confidence: float
    used_chunks: t.List[str]
    provider: str
    model: str


_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?])\s+")


def _split_long_text(text: str, max_len: int) -> t.List[str]:
    parts: t.List[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(length, start + max_len)
        if end < length:
            cut = text.rfind(" ", start, end)
            if cut <= start:
                cut = text.find(" ", end)
            if cut == -1:
                cut = end
            end = cut
        chunk = text[start:end].strip()
        if chunk:
            parts.append(chunk)
        start = max(end, start + 1)
    return parts


def simple_chunk(text: str, max_len: int = 1200) -> t.List[str]:
    if not text:
        return []
    normalized = text.replace("\r\n", "\n")
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", normalized) if p.strip()]
    chunks: t.List[str] = []
    current_sentences: t.List[str] = []
    current_len = 0

    def flush() -> None:
        nonlocal current_sentences, current_len
        if current_sentences:
            chunk_text = " ".join(current_sentences).strip()
            if chunk_text:
                if len(chunk_text) <= max_len:
                    chunks.append(chunk_text)
                else:
                    chunks.extend(_split_long_text(chunk_text, max_len))
        current_sentences = []
        current_len = 0

    for paragraph in paragraphs:
        sentences = _SENTENCE_SPLIT_RE.split(paragraph.strip())
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            sentence_len = len(sentence)
            if current_len + sentence_len + 1 > max_len and current_sentences:
                flush()
            current_sentences.append(sentence)
            current_len += sentence_len + 1
        flush()

    if chunks:
        return chunks
    fallback_text = " ".join(paragraphs).strip()
    return _split_long_text(fallback_text or normalized.strip(), max_len)


class VectorStoreLocal:
    def __init__(self) -> None:
        self.texts: t.List[str] = []
        self.metadatas: t.List[dict] = []
        self.ids: t.List[str] = []
        self.doc_ids: t.List[str] = []
        self.embeddings: t.Optional[np.ndarray] = None
        self._embedder: t.Optional[SentenceTransformer] = None
        self._tfidf: t.Optional[TfidfVectorizer] = None

    def _ensure_embedder(self) -> None:
        if _HAS_ST and self._embedder is None:
            self._embedder = SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

    def _embed(self, texts: t.List[str]) -> np.ndarray:
        if _HAS_ST:
            self._ensure_embedder()
            assert self._embedder is not None
            vecs = self._embedder.encode(texts, convert_to_numpy=True, normalize_embeddings=True)
            return vecs.astype(np.float32)
        if _HAS_SK:
            if self._tfidf is None:
                self._tfidf = TfidfVectorizer(max_features=4096)
                X = self._tfidf.fit_transform(texts)
            else:
                all_texts = self.texts + texts
                self._tfidf = TfidfVectorizer(max_features=4096)
                X = self._tfidf.fit_transform(all_texts)
                self.embeddings = X[: len(self.texts)].toarray().astype(np.float32)
                return X[len(self.texts) :].toarray().astype(np.float32)
            return X.toarray().astype(np.float32)
        def hash_embed(s: str, dim: int = 384) -> np.ndarray:
            vec = np.zeros(dim, dtype=np.float32)
            if not s:
                return vec
            for tok in s.split():
                h = abs(hash(tok)) % dim
                vec[h] += 1.0
            n = np.linalg.norm(vec) + 1e-8
            return (vec / n).astype(np.float32)

        return np.vstack([hash_embed(t) for t in texts])

    def add_documents(self, docs: t.List[IngestDocument]) -> int:
        chunk_texts: t.List[str] = []
        chunk_metas: t.List[dict] = []
        chunk_ids: t.List[str] = []
        doc_ids: t.List[str] = []
        for d in docs:
            base_id = d.id or str(uuid.uuid4())
            for idx, ch in enumerate(simple_chunk(d.text)):
                chunk_texts.append(ch)
                meta = dict(d.meta or {})
                meta.update({"source_id": base_id, "chunk_index": idx})
                chunk_metas.append(meta)
                chunk_ids.append(f"{base_id}#{idx}")
                doc_ids.append(base_id)

        new_emb = self._embed(chunk_texts)
        if self.embeddings is None:
            self.embeddings = new_emb
        else:
            self.embeddings = np.vstack([self.embeddings, new_emb])
        self.texts.extend(chunk_texts)
        self.metadatas.extend(chunk_metas)
        self.ids.extend(chunk_ids)
        self.doc_ids.extend(doc_ids)
        return len(chunk_ids)

    def search(self, query: str, top_k: int = 5) -> t.List[SearchResult]:
        if not self.texts or self.embeddings is None:
            return []
        q = self._embed([query])
        if _HAS_ST or (_HAS_SK and isinstance(self.embeddings, np.ndarray)):
            sims = (q @ self.embeddings.T).reshape(-1) if (_HAS_ST and q.shape[1] == self.embeddings.shape[1]) else cosine_similarity(q, self.embeddings)[0]
        else:
            def cos(a: np.ndarray, b: np.ndarray) -> float:
                return float(np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b) + 1e-8))
            sims = np.array([cos(q[0], v) for v in self.embeddings], dtype=np.float32)

        order = np.argsort(-sims)[:top_k]
        results: t.List[SearchResult] = []
        for i in order:
            results.append(
                SearchResult(
                    doc_id=self.metadatas[i].get("source_id", self.doc_ids[i]),
                    chunk_id=self.ids[i],
                    score=float(sims[i]),
                    text=self.texts[i],
                    meta=self.metadatas[i],
                )
            )
        return results

    def save(self) -> None:
        payload = {
            "texts": self.texts,
            "metadatas": self.metadatas,
            "ids": self.ids,
            "doc_ids": self.doc_ids,
        }
        with open(STORE_PATH, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False)
        if self.embeddings is not None:
            np.save(EMB_PATH, self.embeddings)

    def load(self) -> None:
        if STORE_PATH.exists() and EMB_PATH.exists():
            with open(STORE_PATH, "r", encoding="utf-8") as f:
                payload = json.load(f)
            self.texts = payload.get("texts", [])
            self.metadatas = payload.get("metadatas", [])
            self.ids = payload.get("ids", [])
            self.doc_ids = payload.get("doc_ids", [])
            self.embeddings = np.load(EMB_PATH)


STORE_PATH = DATA_DIR / "store.json"
EMB_PATH = DATA_DIR / "embeddings.npy"
store = VectorStore(STORE_PATH, EMB_PATH)
store.load()

_TXT_MIME_TYPES = {"text/plain"}
_DOC_MIME_TYPES = {
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_DOC_EXT = {".doc", ".docx"}
_TXT_EXT = {".txt"}

app = FastAPI(title="LegalBot AI MVP", version="0.1.0")


def _decode_bytes(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "cp1251", "iso-8859-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_doc_text(content: bytes, ext: str) -> str:
    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        if ext.lower() == ".doc":
            raise HTTPException(
                status_code=400,
                detail="Не удалось прочитать файл .doc. Сохраните документ в формате .docx и попробуйте снова.",
            ) from exc
        raise HTTPException(
            status_code=400,
            detail="Не удалось прочитать файл .docx.",
        ) from exc
    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    return "\n\n".join(paragraphs)


def _extract_text_from_upload_bytes(content: bytes, filename: str, content_type: t.Optional[str]) -> str:
    ext = pathlib.Path(filename or "").suffix.lower()
    if ext in _DOC_EXT or (content_type and content_type.lower() in _DOC_MIME_TYPES):
        return _extract_doc_text(content, ext)
    text_raw = _decode_bytes(content)
    if ext and ext not in _TXT_EXT and content_type and content_type not in _TXT_MIME_TYPES:
        raise HTTPException(status_code=400, detail="Поддерживаются только файлы .doc/.docx и .txt")
    return text_raw


def _sanitize_id(value: str) -> str:
    trimmed = value.strip().replace(" ", "_")
    return trimmed or str(uuid.uuid4())


@app.get("/health")
def health() -> dict:
    return {"status": "ok", "docs": len(store.texts)}


@app.post("/ingest")
def ingest(payload: t.Any = Body(...)) -> dict:
    documents: t.List[IngestDocument] = []
    if isinstance(payload, list):
        if not payload:
            raise HTTPException(status_code=400, detail="documents is empty")
        documents = [IngestDocument(**d) if isinstance(d, dict) else d for d in payload]  # type: ignore
    elif isinstance(payload, dict):
        docs = payload.get("documents")
        if not docs or not isinstance(docs, list):
            raise HTTPException(status_code=400, detail="documents is empty")
        documents = [IngestDocument(**d) if isinstance(d, dict) else d for d in docs]  # type: ignore
    else:
        raise HTTPException(status_code=400, detail="invalid payload format")

    added = store.add_documents(documents)
    store.save()
    return {"added_chunks": added, "total_chunks": len(store.ids)}


@app.post("/ingest/upload")
async def ingest_upload(file: UploadFile = File(...), doc_id: t.Optional[str] = Form(None)) -> dict:
    filename = file.filename or "document"
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Файл пустой")
    text = _extract_text_from_upload_bytes(content, filename, file.content_type).strip()
    if not text:
        raise HTTPException(status_code=400, detail="Не удалось получить текст из файла")

    normalized_id = _sanitize_id(doc_id or pathlib.Path(filename).stem)
    meta = {
        "source_filename": filename,
        "content_type": file.content_type or "",
        "upload_type": "file",
    }
    document = IngestDocument(id=normalized_id, text=text, meta=meta)
    added = store.add_documents([document])
    store.save()
    return {"added_chunks": added, "total_chunks": len(store.ids), "doc_id": normalized_id}


@app.post("/search", response_model=SearchResponse)
def search(req: SearchRequest) -> SearchResponse:
    results = store.search(req.query, top_k=req.top_k)
    return SearchResponse(query=req.query, results=results)


SYSTEM_PROMPT = (
    "Ты — юридический ассистент. Отвечай кратко, строго опираясь на переданный контекст. "
    "Если контекста недостаточно — честно скажи об этом и предложи, какие данные нужны. "
    "Обязательно приводи минимум 2 источника с цитатами: формат [source#chunk] и 1–2 предложения цитаты. "
    "Если релевантных источников меньше двух — отметь это явно и не делай выводов."
)


def _call_llm_openai(prompt: str) -> str:
    if not _HAS_OPENAI:
        return "[OpenAI SDK не установлен]\n\n" + prompt
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return "[OPENAI_API_KEY не задан]\n\n" + prompt
    try:
        from openai import OpenAI  # type: ignore
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=512,
        )
        return resp.choices[0].message.content or ""
    except Exception as e:
        return f"[LLM ошибка: {e}]\n\n" + prompt


def _call_llm_gigachat(prompt: str) -> str:
    token = _gigachat_get_token()
    # Попробуем сначала v2, при 401/403 — фолбэк на v1
    base_url_primary = os.getenv("GIGACHAT_BASE_URL", "https://gigachat.devices.sberbank.ru/api/v2").rstrip('/')
    base_url_fallback = "https://gigachat.devices.sberbank.ru/api/v1"
    model = os.getenv("GIGACHAT_MODEL", "GigaChat:latest")
    verify_ssl = (os.getenv("GIGACHAT_VERIFY", "true").lower() not in ["0", "false", "no"])  # по умолчанию проверяем SSL
    if not token:
        return "[GIGACHAT: не удалось получить токен]\n\n" + prompt
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    x_client_id = os.getenv("GIGACHAT_CLIENT_ID")
    if x_client_id:
        headers["X-Client-ID"] = x_client_id
    headers["RqUID"] = str(uuid.uuid4())
    body = {
        "model": model,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
        "max_tokens": 512,
    }
    try:
        # Первая попытка — v2
        url = f"{base_url_primary}/chat/completions"
        r = requests.post(url, headers=headers, json=body, timeout=60, verify=verify_ssl)
        r.raise_for_status()
        data = r.json()
        return (
            data.get("choices", [{}])[0]
            .get("message", {})
            .get("content", "")
        ) or ""
    except Exception as e:
        # Если был 401/403 — пробуем v1
        try:
            status = getattr(e, 'response', None).status_code if hasattr(e, 'response') and e.response is not None else None
        except Exception:
            status = None
        if status in (401, 403):
            try:
                url_fb = f"{base_url_fallback}/chat/completions"
                r2 = requests.post(url_fb, headers=headers, json=body, timeout=60, verify=verify_ssl)
                r2.raise_for_status()
                data = r2.json()
                return (
                    data.get("choices", [{}])[0]
                    .get("message", {})
                    .get("content", "")
                ) or ""
            except Exception as e2:
                try:
                    body2 = r2.text  # type: ignore[name-defined]
                except Exception:
                    body2 = ""
                suffix2 = f" | body: {body2}" if body2 else ""
                return f"[GIGACHAT ошибка: {e2}{suffix2}]\n\n" + prompt
        # Пытаемся извлечь тело ответа для диагностики
        try:
            body_txt = r.text  # type: ignore[name-defined]
        except Exception:
            body_txt = ""
        suffix = f" | body: {body_txt}" if body_txt else ""
        return f"[GIGACHAT ошибка: {e}{suffix}]\n\n" + prompt


# --------------------------- GigaChat OAuth (token) ---------------------------
_gigachat_token_cache: dict = {"access_token": None, "expires_at": 0.0}


def _gigachat_get_token() -> t.Optional[str]:
    now = time.time()
    cached_token = _gigachat_token_cache.get("access_token")
    expires_at = float(_gigachat_token_cache.get("expires_at") or 0.0)
    if cached_token and now < expires_at - 30:  # 30 сек запас
        return t.cast(str, cached_token)

    # Параметры из окружения
    auth_basic = os.getenv("GIGACHAT_AUTH_BASIC")  # Base64(client_id:secret)
    client_id = os.getenv("GIGACHAT_CLIENT_ID")
    client_secret = os.getenv("GIGACHAT_CLIENT_SECRET")
    scope = os.getenv("GIGACHAT_SCOPE", "GIGACHAT_API_PERS")
    oauth_url = os.getenv("GIGACHAT_OAUTH_URL", "https://ngw.devices.sberbank.ru:9443/api/v2/oauth")
    verify_ssl = (os.getenv("GIGACHAT_VERIFY", "true").lower() not in ["0", "false", "no"])  # общий флаг

    if not auth_basic:
        if client_id and client_secret:
            import base64
            pair = f"{client_id}:{client_secret}".encode()
            auth_basic = base64.b64encode(pair).decode()
        else:
            return None

    headers = {
        "Content-Type": "application/x-www-form-urlencoded",
        "Accept": "application/json",
        "RqUID": str(uuid.uuid4()),
        "Authorization": f"Basic {auth_basic}",
    }
    data = {"scope": scope}
    try:
        resp = requests.post(oauth_url, headers=headers, data=data, timeout=30, verify=verify_ssl)

        print("gigachat token response: ", resp.text)

        resp.raise_for_status()
        j = resp.json()
        access_token = j.get("access_token") or j.get("accessToken")
        # Токен действителен ~30 минут; если пришёл expires_in — используем, иначе 30 мин
        expires_in = float(j.get("expires_in") or 1800)
        if not access_token:
            return None
        _gigachat_token_cache["access_token"] = access_token
        _gigachat_token_cache["expires_at"] = time.time() + expires_in
        return access_token
    except Exception:
        return None


def _call_llm(prompt: str) -> str:
    provider = (os.getenv("LLM_PROVIDER") or "openai").lower()
    if provider == "gigachat":
        return _call_llm_gigachat(prompt)
    # По умолчанию OpenAI; при ошибке попробуем Yandex как фолбэк
    out = _call_llm_openai(prompt)
    return out


@app.post("/answer", response_model=AnswerResponse)
def answer(req: AnswerRequest) -> AnswerResponse:
    results_raw = store.search(req.query, top_k=req.top_k)
    # Нормализуем к типу SearchResult
    results: t.List[SearchResult] = [
        (SearchResult(**r) if isinstance(r, dict) else r)  # type: ignore
        for r in results_raw
    ]
    if not results:
        raise HTTPException(status_code=404, detail="Нет релевантных фрагментов. Добавьте документы через /ingest.")

    citations = []
    for r in results:
        citations.append(
            {
                "id": r.chunk_id,
                "score": round(r.score, 4),
                "meta": r.meta,
                "quote": (r.text[:400] + ("…" if len(r.text) > 400 else "")),
            }
        )

    context_blocks = []
    for r in results:
        src = r.meta.get("source_id", r.doc_id)
        context_blocks.append(f"[{src}/{r.chunk_id}]\n{r.text}")
    context = "\n\n---\n\n".join(context_blocks)

    user_prompt = (
        f"Вопрос: {req.query}\n\n"
        f"Контекст (фрагменты):\n\n{context}\n\n"
        "Ответ: кратко, по пунктам; укажи источники в виде [source#chunk]."
    )

    # Определим провайдера/модель для ответа
    provider_env = (os.getenv("LLM_PROVIDER") or "openai").lower()
    if provider_env == "gigachat":
        model_name = os.getenv("GIGACHAT_MODEL", "GigaChat:latest")
    else:
        model_name = "gpt-4o-mini"
    answer_text = _call_llm(user_prompt)
    used_ids = [r.chunk_id for r in results]
    confidence = float(np.mean([r.score for r in results])) if results else 0.0
    return AnswerResponse(answer=answer_text, citations=citations, confidence=confidence, used_chunks=used_ids, provider=provider_env, model=model_name)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app.main:app", host="0.0.0.0", port=int(os.getenv("PORT", "8000")), reload=False)
